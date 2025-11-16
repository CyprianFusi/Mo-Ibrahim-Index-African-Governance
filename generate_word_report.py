"""
Comprehensive MS Word Report Generator for Ibrahim Index of African Governance (IIAG)
Generates a detailed Word document with cover page, analysis, and embedded visualizations
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import warnings
from math import pi
warnings.filterwarnings('ignore')

# Set professional styling for charts
plt.style.use('seaborn-v0_8-whitegrid')
sns.set_palette("husl")

# Load data
print("Loading data...")
data_path = Path('data/csv-files')
composite_scores = pd.read_csv(data_path / '2024 IIAG_Composite Scores.csv', encoding='utf-8-sig')

# Clean data
composite_scores = composite_scores.replace('.', np.nan)
for col in composite_scores.columns[3:]:
    composite_scores[col] = pd.to_numeric(composite_scores[col], errors='coerce')

# Define categories
main_categories = [
    'SECURITY & RULE OF LAW',
    'PARTICIPATION, RIGHTS & INCLUSION',
    'FOUNDATIONS FOR ECONOMIC OPPORTUNITY',
    'HUMAN DEVELOPMENT'
]

subcategories = {
    'SECURITY & RULE OF LAW': ['SECURITY & SAFETY', 'RULE OF LAW & JUSTICE', 'ACCOUNTABILITY & TRANSPARENCY', 'ANTI-CORRUPTION'],
    'PARTICIPATION, RIGHTS & INCLUSION': ['PARTICIPATION', 'RIGHTS', 'INCLUSION & EQUALITY', "WOMEN'S EQUALITY"],
    'FOUNDATIONS FOR ECONOMIC OPPORTUNITY': ['PUBLIC ADMINISTRATION', 'BUSINESS & LABOUR ENVIRONMENT', 'INFRASTRUCTURE', 'RURAL ECONOMY'],
    'HUMAN DEVELOPMENT': ['HEALTH', 'EDUCATION', 'SOCIAL PROTECTION & WELFARE', 'SUSTAINABLE ENVIRONMENT']
}

# Regional groupings
regional_groups = {
    'North Africa': ['Algeria', 'Egypt', 'Libya', 'Morocco', 'Tunisia'],
    'West Africa': ['Benin', 'Burkina Faso', 'Cape Verde', 'Côte d\'Ivoire', 'Gambia', 'Ghana', 'Guinea',
                    'Guinea-Bissau', 'Liberia', 'Mali', 'Mauritania', 'Niger', 'Nigeria', 'Senegal',
                    'Sierra Leone', 'Togo'],
    'East Africa': ['Burundi', 'Comoros', 'Djibouti', 'Eritrea', 'Ethiopia', 'Kenya', 'Madagascar',
                    'Mauritius', 'Rwanda', 'Seychelles', 'Somalia', 'South Sudan', 'Sudan', 'Tanzania', 'Uganda'],
    'Central Africa': ['Cameroon', 'Central African Republic', 'Chad', 'Congo', 'DR Congo', 'Equatorial Guinea',
                       'Gabon', 'São Tomé and Príncipe'],
    'Southern Africa': ['Angola', 'Botswana', 'Eswatini', 'Lesotho', 'Malawi', 'Mozambique', 'Namibia',
                        'South Africa', 'Zambia', 'Zimbabwe']
}

def get_region(country):
    for region, countries in regional_groups.items():
        if country in countries:
            return region
    return 'Other'

composite_scores['Region'] = composite_scores['Country'].apply(get_region)

latest_year = composite_scores['Year'].max()
earliest_year = composite_scores['Year'].min()
latest_data = composite_scores[composite_scores['Year'] == latest_year].copy()

# Calculate changes
def calculate_change(df, country, start_year, end_year):
    start_score = df[(df['Country'] == country) & (df['Year'] == start_year)]['OVERALL GOVERNANCE'].values
    end_score = df[(df['Country'] == country) & (df['Year'] == end_year)]['OVERALL GOVERNANCE'].values
    if len(start_score) > 0 and len(end_score) > 0:
        return end_score[0] - start_score[0]
    return np.nan

countries_list = composite_scores['Country'].unique()
changes = []
for country in countries_list:
    change = calculate_change(composite_scores, country, earliest_year, latest_year)
    if not np.isnan(change):
        changes.append({'Country': country, 'Change': change})

changes_df = pd.DataFrame(changes).sort_values('Change', ascending=False)

# Create temporary directory for charts
temp_dir = Path('temp_charts')
temp_dir.mkdir(exist_ok=True)

print("Generating charts for Word document...")

# Generate all charts
chart_files = []

# Chart 1: Distribution
print("  Creating distribution chart...")
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

ax1.hist(latest_data['OVERALL GOVERNANCE'].dropna(), bins=20, edgecolor='black', alpha=0.7, color='#3498db')
ax1.axvline(latest_data['OVERALL GOVERNANCE'].mean(), color='red', linestyle='--', linewidth=2,
            label=f'Mean: {latest_data["OVERALL GOVERNANCE"].mean():.1f}')
ax1.axvline(latest_data['OVERALL GOVERNANCE'].median(), color='green', linestyle='--', linewidth=2,
            label=f'Median: {latest_data["OVERALL GOVERNANCE"].median():.1f}')
ax1.set_xlabel('Overall Governance Score', fontweight='bold', fontsize=11)
ax1.set_ylabel('Number of Countries', fontweight='bold', fontsize=11)
ax1.set_title(f'Distribution of Governance Scores ({latest_year})', fontweight='bold', fontsize=13)
ax1.legend()
ax1.grid(True, alpha=0.3)

regional_stats = latest_data.groupby('Region')['OVERALL GOVERNANCE'].mean().sort_values(ascending=False)
regional_stats = regional_stats[regional_stats.index != 'Other']
region_data = [latest_data[latest_data['Region'] == region]['OVERALL GOVERNANCE'].dropna()
               for region in regional_stats.index]

bp = ax2.boxplot(region_data, labels=regional_stats.index, patch_artist=True)
for patch in bp['boxes']:
    patch.set_facecolor('#3498db')
    patch.set_alpha(0.7)
ax2.set_ylabel('Overall Governance Score', fontweight='bold', fontsize=11)
ax2.set_xlabel('Region', fontweight='bold', fontsize=11)
ax2.set_title(f'Regional Governance Comparison ({latest_year})', fontweight='bold', fontsize=13)
ax2.tick_params(axis='x', rotation=45)
ax2.grid(True, alpha=0.3, axis='y')

plt.tight_layout()
chart1 = temp_dir / 'chart1_distribution.png'
plt.savefig(chart1, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart1)

# Chart 2: Top/Bottom Countries
print("  Creating top/bottom countries chart...")
fig, ax = plt.subplots(figsize=(12, 10))

top_bottom = pd.concat([
    latest_data.nlargest(15, 'OVERALL GOVERNANCE'),
    latest_data.nsmallest(15, 'OVERALL GOVERNANCE')
]).sort_values('OVERALL GOVERNANCE')

colors = ['#e74c3c' if x < 50 else '#f39c12' if x < 60 else '#2ecc71'
          for x in top_bottom['OVERALL GOVERNANCE']]
bars = ax.barh(range(len(top_bottom)), top_bottom['OVERALL GOVERNANCE'],
               color=colors, edgecolor='black', linewidth=0.7)

ax.set_yticks(range(len(top_bottom)))
ax.set_yticklabels(top_bottom['Country'], fontsize=9)
ax.set_xlabel('Overall Governance Score', fontweight='bold', fontsize=11)
ax.set_title(f'Top and Bottom 15 Countries - Overall Governance ({latest_year})',
             fontweight='bold', fontsize=14, pad=15)
ax.grid(True, alpha=0.3, axis='x')
ax.axvline(50, color='black', linestyle='--', linewidth=1, alpha=0.5)

for i, (bar, val) in enumerate(zip(bars, top_bottom['OVERALL GOVERNANCE'])):
    ax.text(val + 1, i, f'{val:.1f}', va='center', fontweight='bold', fontsize=8)

plt.tight_layout()
chart2 = temp_dir / 'chart2_top_bottom.png'
plt.savefig(chart2, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart2)

# Chart 3: Temporal Trends
print("  Creating temporal trends chart...")
yearly_avg = composite_scores.groupby('Year')['OVERALL GOVERNANCE'].mean()
yearly_categories = composite_scores.groupby('Year')[main_categories].mean()

fig, ax = plt.subplots(figsize=(14, 7))

ax.plot(yearly_avg.index, yearly_avg.values, marker='o', linewidth=3, markersize=10,
        label='Overall Governance', color='#2c3e50')

colors_cat = ['#e74c3c', '#3498db', '#2ecc71', '#f39c12']
for idx, cat in enumerate(main_categories):
    ax.plot(yearly_categories.index, yearly_categories[cat], marker='s', linewidth=2.5,
            markersize=7, label=cat, alpha=0.8, color=colors_cat[idx])

ax.set_xlabel('Year', fontweight='bold', fontsize=11)
ax.set_ylabel('Average Governance Score', fontweight='bold', fontsize=11)
ax.set_title(f'Continental Governance Trends ({earliest_year}-{latest_year})',
             fontweight='bold', fontsize=14, pad=15)
ax.legend(loc='best', framealpha=0.95, fontsize=9)
ax.grid(True, alpha=0.3)

plt.tight_layout()
chart3 = temp_dir / 'chart3_trends.png'
plt.savefig(chart3, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart3)

# Chart 4: Category Heatmap
print("  Creating category heatmap...")
top20_data = latest_data.nlargest(20, 'OVERALL GOVERNANCE')
heatmap_data = top20_data[['Country'] + main_categories].set_index('Country')

fig, ax = plt.subplots(figsize=(11, 9))
sns.heatmap(heatmap_data, annot=True, fmt='.1f', cmap='RdYlGn', center=50,
            linewidths=0.5, cbar_kws={'label': 'Score'}, ax=ax, vmin=30, vmax=90)
ax.set_title(f'Category Performance - Top 20 Countries ({latest_year})',
             fontweight='bold', fontsize=13, pad=15)
ax.set_xlabel('')
ax.set_ylabel('Country', fontweight='bold', fontsize=11)

plt.tight_layout()
chart4 = temp_dir / 'chart4_heatmap.png'
plt.savefig(chart4, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart4)

# Chart 5: Governance Change
print("  Creating governance change chart...")
fig, ax = plt.subplots(figsize=(12, 11))

colors_change = ['#27ae60' if x > 0 else '#e74c3c' for x in changes_df['Change']]
bars = ax.barh(range(len(changes_df)), changes_df['Change'],
               color=colors_change, edgecolor='black', linewidth=0.5)

ax.set_yticks(range(len(changes_df)))
ax.set_yticklabels(changes_df['Country'], fontsize=7)
ax.set_xlabel(f'Change in Overall Governance Score ({earliest_year}-{latest_year})',
              fontweight='bold', fontsize=11)
ax.set_title(f'Governance Change: All Countries ({earliest_year}-{latest_year})',
             fontweight='bold', fontsize=14, pad=15)
ax.axvline(0, color='black', linewidth=2)
ax.grid(True, alpha=0.3, axis='x')

for i in list(range(10)) + list(range(len(changes_df)-10, len(changes_df))):
    val = changes_df['Change'].iloc[i]
    ax.text(val + (0.5 if val > 0 else -0.5), i, f'{val:.1f}', va='center',
            ha='left' if val > 0 else 'right', fontweight='bold', fontsize=7)

plt.tight_layout()
chart5 = temp_dir / 'chart5_change.png'
plt.savefig(chart5, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart5)

# Chart 6: Regional Comparison
print("  Creating regional comparison chart...")
regional_means = latest_data.groupby('Region')[['OVERALL GOVERNANCE'] + main_categories].mean()
regional_means = regional_means[regional_means.index != 'Other'].sort_values('OVERALL GOVERNANCE', ascending=False)

fig, ax = plt.subplots(figsize=(14, 7))
x = np.arange(len(regional_means))
width = 0.15

colors_bar = ['#2c3e50', '#e74c3c', '#3498db', '#2ecc71', '#f39c12']
for i, cat in enumerate(['OVERALL GOVERNANCE'] + main_categories):
    offset = width * (i - 2)
    ax.bar(x + offset, regional_means[cat], width, label=cat, color=colors_bar[i])

ax.set_xlabel('Region', fontweight='bold', fontsize=11)
ax.set_ylabel('Average Score', fontweight='bold', fontsize=11)
ax.set_title(f'Regional Performance Across Categories ({latest_year})',
             fontweight='bold', fontsize=14, pad=15)
ax.set_xticks(x)
ax.set_xticklabels(regional_means.index, rotation=20, ha='right')
ax.legend(loc='best', framealpha=0.95, fontsize=9)
ax.grid(True, alpha=0.3, axis='y')

plt.tight_layout()
chart6 = temp_dir / 'chart6_regional.png'
plt.savefig(chart6, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart6)

# Chart 7: Correlation
print("  Creating correlation chart...")
fig, axes = plt.subplots(2, 2, figsize=(14, 12))
axes = axes.flatten()

for idx, cat in enumerate(main_categories):
    ax = axes[idx]
    plot_data = latest_data[['OVERALL GOVERNANCE', cat]].dropna()

    ax.scatter(plot_data[cat], plot_data['OVERALL GOVERNANCE'],
               alpha=0.6, s=80, edgecolors='black', linewidth=0.5, color='#3498db')

    z = np.polyfit(plot_data[cat], plot_data['OVERALL GOVERNANCE'], 1)
    p = np.poly1d(z)
    ax.plot(plot_data[cat].sort_values(), p(plot_data[cat].sort_values()),
            "r--", linewidth=2.5, alpha=0.8)

    corr = plot_data[cat].corr(plot_data['OVERALL GOVERNANCE'])

    ax.set_xlabel(cat, fontweight='bold', fontsize=10)
    ax.set_ylabel('Overall Governance Score', fontweight='bold', fontsize=10)
    ax.set_title(f'{cat}\n(Correlation: {corr:.3f})', fontweight='bold', fontsize=11)
    ax.grid(True, alpha=0.3)

plt.suptitle(f'Category vs Overall Governance Correlation ({latest_year})',
             fontweight='bold', fontsize=14, y=0.995)
plt.tight_layout()
chart7 = temp_dir / 'chart7_correlation.png'
plt.savefig(chart7, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart7)

# Chart 8: Top/Bottom Trends
print("  Creating top/bottom trends chart...")
top5 = latest_data.nlargest(5, 'OVERALL GOVERNANCE')['Country'].values
bottom5 = latest_data.nsmallest(5, 'OVERALL GOVERNANCE')['Country'].values

fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))

for country in top5:
    country_data = composite_scores[composite_scores['Country'] == country]
    ax1.plot(country_data['Year'], country_data['OVERALL GOVERNANCE'],
             marker='o', linewidth=2.5, label=country, markersize=7)

ax1.set_xlabel('Year', fontweight='bold', fontsize=11)
ax1.set_ylabel('Overall Governance Score', fontweight='bold', fontsize=11)
ax1.set_title(f'Top 5 Performing Countries - Trends ({earliest_year}-{latest_year})',
              fontweight='bold', fontsize=12)
ax1.legend(loc='best', framealpha=0.95)
ax1.grid(True, alpha=0.3)

for country in bottom5:
    country_data = composite_scores[composite_scores['Country'] == country]
    ax2.plot(country_data['Year'], country_data['OVERALL GOVERNANCE'],
             marker='s', linewidth=2.5, label=country, markersize=7)

ax2.set_xlabel('Year', fontweight='bold', fontsize=11)
ax2.set_ylabel('Overall Governance Score', fontweight='bold', fontsize=11)
ax2.set_title(f'Bottom 5 Performing Countries - Trends ({earliest_year}-{latest_year})',
              fontweight='bold', fontsize=12)
ax2.legend(loc='best', framealpha=0.95)
ax2.grid(True, alpha=0.3)

plt.tight_layout()
chart8 = temp_dir / 'chart8_top_bottom_trends.png'
plt.savefig(chart8, dpi=300, bbox_inches='tight')
plt.close()
chart_files.append(chart8)

# Create Word Document
print("\nCreating Word document...")
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ========== COVER PAGE ==========
print("  Adding cover page...")
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("IBRAHIM INDEX OF AFRICAN GOVERNANCE\n\n")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("COMPREHENSIVE ANALYTICAL REPORT")
run.font.size = Pt(22)
run.font.bold = True

doc.add_paragraph("\n" * 2)

period = doc.add_paragraph()
period.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = period.add_run(f"Analysis Period: {earliest_year}-{latest_year}")
run.font.size = Pt(14)
run.font.italic = True

doc.add_paragraph("\n" * 2)

# Key stats
stats = doc.add_paragraph()
stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = stats.add_run("Dataset Overview\n\n")
run.font.size = Pt(14)
run.font.bold = True

stats_content = doc.add_paragraph()
stats_content.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = stats_content.add_run(
    f"Countries Analyzed: {composite_scores['Country'].nunique()}\n"
    f"Years Covered: {earliest_year} - {latest_year}\n"
    f"Total Observations: {len(composite_scores):,}\n"
    f"Governance Categories: 4 Main Categories, 16 Subcategories"
)
run.font.size = Pt(12)

doc.add_paragraph("\n" * 4)

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\nMo Ibrahim Foundation")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ========== EXECUTIVE SUMMARY ==========
print("  Adding executive summary...")
heading = doc.add_heading('EXECUTIVE SUMMARY', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph(
    f"This report presents a comprehensive analysis of governance performance across African nations "
    f"from {earliest_year} to {latest_year}, based on the Ibrahim Index of African Governance (IIAG)."
)

doc.add_heading('KEY FINDINGS', 1)

doc.add_heading(f'1. Overall Governance Landscape ({latest_year})', 2)
para = doc.add_paragraph()
para.add_run(f"• Continental Mean Score: {latest_data['OVERALL GOVERNANCE'].mean():.1f}/100\n")
para.add_run(f"• Median Score: {latest_data['OVERALL GOVERNANCE'].median():.1f}/100\n")
para.add_run(f"• Score Range: {latest_data['OVERALL GOVERNANCE'].min():.1f} to {latest_data['OVERALL GOVERNANCE'].max():.1f}\n")
para.add_run(f"• Standard Deviation: {latest_data['OVERALL GOVERNANCE'].std():.1f}")

doc.add_heading(f'2. Top Performers ({latest_year})', 2)
doc.add_paragraph("The top five countries demonstrate exceptional governance:")
for _, row in latest_data.nlargest(5, 'OVERALL GOVERNANCE').iterrows():
    doc.add_paragraph(f"{row['Country']}: {row['OVERALL GOVERNANCE']:.1f}", style='List Bullet')

doc.add_heading(f'3. Significant Improvers ({earliest_year}-{latest_year})', 2)
doc.add_paragraph("Countries showing the greatest governance improvements:")
for _, row in changes_df.head(5).iterrows():
    doc.add_paragraph(f"{row['Country']}: +{row['Change']:.1f} points", style='List Bullet')

doc.add_heading('4. Regional Patterns', 2)
regional_summary = latest_data.groupby('Region')['OVERALL GOVERNANCE'].mean().sort_values(ascending=False)
for region, score in regional_summary.head(5).items():
    if region != 'Other':
        doc.add_paragraph(f"{region}: {score:.1f}", style='List Bullet')

doc.add_heading('5. Category Performance', 2)
doc.add_paragraph("Average scores across main governance categories:")
for cat in main_categories:
    doc.add_paragraph(f"{cat}: {latest_data[cat].mean():.1f}", style='List Bullet')

doc.add_heading('METHODOLOGY', 2)
doc.add_paragraph(
    "The IIAG assesses governance across four main categories: Security & Rule of Law, Participation, "
    "Rights & Inclusion, Foundations for Economic Opportunity, and Human Development. Each category "
    "comprises multiple subcategories, creating a comprehensive governance assessment framework."
)

doc.add_page_break()

# ========== VISUALIZATIONS AND ANALYSIS ==========
print("  Adding visualizations...")

doc.add_heading('1. GOVERNANCE DISTRIBUTION AND REGIONAL COMPARISON', 1)
doc.add_paragraph(
    f"Figure 1 illustrates the distribution of governance scores across all African countries in {latest_year}. "
    f"The histogram reveals the central tendency and spread of governance performance, while the boxplot "
    f"compares regional variations. The continental mean of {latest_data['OVERALL GOVERNANCE'].mean():.1f} "
    f"and median of {latest_data['OVERALL GOVERNANCE'].median():.1f} indicate the typical governance level."
)
doc.add_picture(str(chart_files[0]), width=Inches(6.5))

doc.add_page_break()

doc.add_heading('2. TOP AND BOTTOM PERFORMING COUNTRIES', 1)
doc.add_paragraph(
    f"Figure 2 presents the top 15 and bottom 15 countries by overall governance score in {latest_year}. "
    f"The color coding highlights different performance levels: green indicates strong governance (60+), "
    f"orange represents moderate performance (50-60), and red signifies governance challenges (below 50). "
    f"This visualization reveals the significant governance gap across the continent."
)
doc.add_picture(str(chart_files[1]), width=Inches(6.5))

doc.add_page_break()

doc.add_heading('3. CONTINENTAL GOVERNANCE TRENDS', 1)
doc.add_paragraph(
    f"Figure 3 tracks the evolution of overall governance and its four main categories from {earliest_year} "
    f"to {latest_year}. This temporal analysis reveals long-term trends and patterns in African governance. "
    f"The data shows that overall governance has "
    f"{'improved' if latest_data['OVERALL GOVERNANCE'].mean() > composite_scores[composite_scores['Year']==earliest_year]['OVERALL GOVERNANCE'].mean() else 'declined'} "
    f"over the analysis period."
)
doc.add_picture(str(chart_files[2]), width=Inches(6.5))

doc.add_page_break()

doc.add_heading('4. CATEGORY PERFORMANCE HEATMAP', 1)
doc.add_paragraph(
    f"Figure 4 displays a detailed heatmap of category performance for the top 20 countries in {latest_year}. "
    f"This visualization enables identification of governance strengths and weaknesses across different "
    f"dimensions. Green cells indicate strong performance, yellow represents moderate scores, and red "
    f"highlights areas requiring improvement."
)
doc.add_picture(str(chart_files[3]), width=Inches(6.5))

doc.add_page_break()

doc.add_heading('5. GOVERNANCE CHANGE ANALYSIS', 1)
improvers = changes_df[changes_df['Change'] > 0]
decliners = changes_df[changes_df['Change'] < 0]
doc.add_paragraph(
    f"Figure 5 presents governance changes for all countries from {earliest_year} to {latest_year}. "
    f"Green bars indicate improvement, while red bars show decline. Of the {len(changes_df)} countries "
    f"analyzed, {len(improvers)} ({len(improvers)/len(changes_df)*100:.1f}%) showed improvement, while "
    f"{len(decliners)} ({len(decliners)/len(changes_df)*100:.1f}%) experienced decline."
)
doc.add_picture(str(chart_files[4]), width=Inches(6.5))

doc.add_page_break()

doc.add_heading('6. REGIONAL PERFORMANCE COMPARISON', 1)
doc.add_paragraph(
    f"Figure 6 compares regional performance across all governance categories in {latest_year}. "
    f"This clustered bar chart reveals regional patterns and highlights which regions excel in specific "
    f"governance dimensions. The comparison facilitates understanding of regional strengths and areas "
    f"for targeted improvement."
)
doc.add_picture(str(chart_files[5]), width=Inches(6.5))

doc.add_page_break()

doc.add_heading('7. CATEGORY CORRELATION ANALYSIS', 1)
doc.add_paragraph(
    f"Figure 7 presents scatter plots examining the relationship between each main category and overall "
    f"governance scores in {latest_year}. The correlation coefficients quantify the strength of these "
    f"relationships, while trend lines visualize the associations. These correlations indicate which "
    f"governance dimensions most strongly influence overall performance."
)
doc.add_picture(str(chart_files[6]), width=Inches(6.5))

doc.add_page_break()

doc.add_heading('8. TOP AND BOTTOM COUNTRY TRENDS', 1)
doc.add_paragraph(
    f"Figure 8 tracks the governance trajectories of the top 5 and bottom 5 performing countries from "
    f"{earliest_year} to {latest_year}. These trend lines reveal whether high performers maintain their "
    f"positions, whether struggling countries show improvement, and the stability or volatility of "
    f"governance scores over time."
)
doc.add_picture(str(chart_files[7]), width=Inches(6.5))

doc.add_page_break()

# ========== DETAILED ANALYSIS ==========
print("  Adding detailed analysis...")
doc.add_heading('DETAILED ANALYSIS AND INSIGHTS', 1)

doc.add_heading(f'1. Temporal Dynamics ({earliest_year}-{latest_year})', 2)
doc.add_paragraph("Continental Trend:")
early_mean = composite_scores[composite_scores['Year']==earliest_year]['OVERALL GOVERNANCE'].mean()
para = doc.add_paragraph()
para.add_run(f"• Overall governance score changed from {early_mean:.1f} ({earliest_year}) to {latest_data['OVERALL GOVERNANCE'].mean():.1f} ({latest_year})\n")
para.add_run(f"• Net change: {latest_data['OVERALL GOVERNANCE'].mean() - early_mean:.2f} points\n")
para.add_run(f"• Countries improving: {len(improvers)} ({len(improvers)/len(changes_df)*100:.1f}%)\n")
para.add_run(f"• Countries declining: {len(decliners)} ({len(decliners)/len(changes_df)*100:.1f}%)")

doc.add_heading(f'2. Category-Specific Insights ({latest_year})', 2)
cat_means = [(cat, latest_data[cat].mean()) for cat in main_categories]
strongest = max(cat_means, key=lambda x: x[1])
weakest = min(cat_means, key=lambda x: x[1])

para = doc.add_paragraph()
para.add_run(f"Strongest Category (Continental Average):\n")
para.add_run(f"• {strongest[0]}: {strongest[1]:.1f}\n\n")
para.add_run(f"Weakest Category (Continental Average):\n")
para.add_run(f"• {weakest[0]}: {weakest[1]:.1f}\n\n")
para.add_run(f"Category Variability (Standard Deviation):\n")
for cat in main_categories:
    para.add_run(f"• {cat}: {latest_data[cat].std():.1f}\n")

doc.add_heading('3. Regional Performance Dynamics', 2)
regional_changes = []
for region in regional_groups.keys():
    region_countries = [c for c in countries_list if get_region(c) == region]
    region_change = changes_df[changes_df['Country'].isin(region_countries)]['Change'].mean()
    if not np.isnan(region_change):
        regional_changes.append((region, region_change))
regional_changes.sort(key=lambda x: x[1], reverse=True)

doc.add_paragraph(f"Regional Rankings by Average Change ({earliest_year}-{latest_year}):")
for i, (region, change) in enumerate(regional_changes[:5], 1):
    doc.add_paragraph(f"{i}. {region}: {change:+.2f} points", style='List Number')

doc.add_heading('4. Notable Patterns and Observations', 2)
corr_security = latest_data[['SECURITY & RULE OF LAW', 'OVERALL GOVERNANCE']].corr().iloc[0,1]
high_performers = len(latest_data[latest_data['OVERALL GOVERNANCE'] >= 60])
low_performers = len(latest_data[latest_data['OVERALL GOVERNANCE'] < 50])

para = doc.add_paragraph()
para.add_run(f"• The correlation between 'Security & Rule of Law' and 'Overall Governance' is {corr_security:.3f}, "
             f"indicating {'strong' if abs(corr_security) > 0.8 else 'moderate'} relationship\n\n")
para.add_run(f"• {high_performers} countries ({high_performers/len(latest_data)*100:.1f}%) achieved governance "
             f"scores above 60/100\n\n")
para.add_run(f"• {low_performers} countries ({low_performers/len(latest_data)*100:.1f}%) scored below 50/100, "
             f"indicating significant governance challenges")

doc.add_heading('5. Methodology Notes', 2)
doc.add_paragraph(
    "The Ibrahim Index of African Governance (IIAG) provides a comprehensive assessment framework:\n"
    "• Covers 54 African countries\n"
    "• Uses approximately 100 indicators from approximately 30 independent sources\n"
    "• Scores range from 0-100 (higher is better)\n"
    "• Four main categories with equal weighting\n"
    "• Annual updates reflecting latest available data"
)

doc.add_page_break()

# ========== REFERENCES ==========
print("  Adding references...")
doc.add_heading('REFERENCES', 1)

refs = [
    f"Mo Ibrahim Foundation (2024). Ibrahim Index of African Governance (IIAG) 2024. Available at: https://mo.ibrahim.foundation/iiag [Accessed: {datetime.now().strftime('%d %B %Y')}].",
    "\nMo Ibrahim Foundation (2024). '2024 IIAG Composite Scores Dataset'. Mo Ibrahim Foundation Data Portal.",
    "\nMo Ibrahim Foundation (2023). 2023 Ibrahim Index of African Governance: Index Report. London: Mo Ibrahim Foundation.",
    "\nAfrican Development Bank (2023). African Economic Outlook 2023. Abidjan: African Development Bank Group.",
    "\nUnited Nations Development Programme (2023). Human Development Report 2023. New York: UNDP.",
    "\nWorld Bank (2023). Worldwide Governance Indicators 2023. Washington DC: World Bank Group."
]

for ref in refs:
    doc.add_paragraph(ref)

doc.add_heading('ABOUT THE IBRAHIM INDEX OF AFRICAN GOVERNANCE (IIAG)', 2)
doc.add_paragraph(
    "The IIAG is the most comprehensive assessment of African governance, providing an annual "
    "statistical measure of governance performance in every African country. Launched in 2007 "
    "by the Mo Ibrahim Foundation, it covers all 54 African countries and is based on data from "
    "over 30 independent African and global institutions."
)

doc.add_paragraph(
    "The Index measures governance performance across four main categories: Security & Rule of Law, "
    "Participation, Rights & Inclusion, Foundations for Economic Opportunity, and Human Development. "
    "Each category comprises subcategories and indicators that assess different dimensions of "
    "governance, from civil liberties to infrastructure development."
)

doc.add_heading('METHODOLOGY', 2)
para = doc.add_paragraph()
para.add_run("Data Sources: ").bold = True
para.add_run("The IIAG uses approximately 100 indicators from reputable independent sources "
             "including international organizations, research institutes, and African institutions.\n\n")
para.add_run("Scoring: ").bold = True
para.add_run("All indicators are converted to a 0-100 scale where 100 represents the best "
             "possible outcome. Country scores are calculated as weighted averages of indicator values.\n\n")
para.add_run("Coverage: ").bold = True
para.add_run(f"The analysis in this report covers {earliest_year}-{latest_year}, examining trends "
             f"across {len(countries_list)} countries and providing insights into governance performance at "
             "continental, regional, and national levels.")

doc.add_heading('ACKNOWLEDGMENTS', 2)
doc.add_paragraph(
    "This report utilizes data provided by the Mo Ibrahim Foundation under their open data "
    "policy. The analysis, interpretations, and conclusions presented are those of the report "
    "author and do not necessarily reflect the views of the Mo Ibrahim Foundation."
)

# Save document
doc_filename = f'IIAG_Comprehensive_Report_{latest_year}.docx'
doc.save(doc_filename)

# Clean up temporary files
print("\nCleaning up temporary files...")
import shutil
shutil.rmtree(temp_dir)

print(f"\n{'='*80}")
print("WORD REPORT GENERATION COMPLETE!")
print(f"{'='*80}")
print(f"\nReport saved as: {doc_filename}")
print(f"Pages: ~15-20")
print("\nThe Word document includes:")
print("  • Professional cover page")
print("  • Executive summary")
print("  • 8 high-quality embedded charts")
print("  • Detailed analysis sections")
print("  • Harvard-style references")
print("  • Fully editable formatting")
print(f"\n{'='*80}")
